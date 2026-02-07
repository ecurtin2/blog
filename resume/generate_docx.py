#!/usr/bin/env python3
"""Generate styled DOCX resume from profile.toml with blog theme colors."""

import sys
try:
    import tomllib  # Python 3.11+
except ModuleNotFoundError:  # pragma: no cover
    try:
        import tomli as tomllib  # Python <=3.10 backport
    except ModuleNotFoundError as e:  # pragma: no cover
        raise ModuleNotFoundError(
            "Missing TOML parser. Use Python 3.11+ (tomllib) or install 'tomli'."
        ) from e
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Blog theme colors
TEAL = RGBColor(0x03, 0x71, 0x71)
GREEN = RGBColor(0x2A, 0x7F, 0x62)
DARKGREEN = RGBColor(0x03, 0x31, 0x2E)
BLUE = RGBColor(0x00, 0x66, 0xBB)
FG = RGBColor(0x1D, 0x1D, 0x1D)
LIGHT_GRAY = RGBColor(0x66, 0x66, 0x66)


def add_hyperlink(paragraph, text, url):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0066BB")
    rPr.append(color)
    
    new_run.append(rPr)
    
    # Use w:t element for text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_paragraph_spacing(paragraph, before=0, after=0):
    """Set paragraph spacing."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    pPr.append(spacing)


def add_bottom_border(paragraph, color="2A7F62"):
    """Add bottom border to paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def included_in(item, target: str) -> bool:
    """Return True if item is explicitly included for the target output."""
    if not isinstance(item, dict):
        return False
    return target in item.get("included_in", [])


def filter_included(items, target: str):
    return [x for x in items if included_in(x, target)]


def main():
    if len(sys.argv) < 2:
        print("Usage: generate_docx.py <output.docx>", file=sys.stderr)
        sys.exit(1)
    
    output_path = sys.argv[1]
    profile_path = Path(__file__).parent.parent / "data" / "profile.toml"
    
    with open(profile_path, "rb") as f:
        profile = tomllib.load(f)

    p = profile["personal"]
    bio = profile["bio"]
    skills = profile["skills"]
    target = "resume"

    doc = Document()
    
    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = FG

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(p["name"])
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = TEAL
    set_paragraph_spacing(name_para, after=0)

    # Tagline
    tagline_para = doc.add_paragraph()
    tagline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline_run = tagline_para.add_run(bio["tagline"])
    tagline_run.italic = True
    tagline_run.font.color.rgb = LIGHT_GRAY
    set_paragraph_spacing(tagline_para, before=0, after=60)

    # Contact line
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f'{p["location"]} | {p["email"]} | {p["phone"]} | github.com/ecurtin2'
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(10)
    set_paragraph_spacing(contact_para, after=120)

    # Skills section
    skills_header = doc.add_paragraph()
    skills_run = skills_header.add_run("SKILLS")
    skills_run.bold = True
    skills_run.font.size = Pt(12)
    skills_run.font.color.rgb = GREEN
    add_bottom_border(skills_header)
    set_paragraph_spacing(skills_header, before=200, after=100)

    for group in skills.get("groups", []):
        items = [it.get("name", "") for it in filter_included(group.get("items", []), target) if isinstance(it, dict) and it.get("name")]
        if not items:
            continue
        skill_para = doc.add_paragraph()
        label_run = skill_para.add_run(f'{group.get("label", "Skills")}: ')
        label_run.bold = True
        label_run.font.color.rgb = DARKGREEN
        skill_para.add_run(", ".join(items))
        set_paragraph_spacing(skill_para, after=40)

    # Experience section
    exp_header = doc.add_paragraph()
    exp_run = exp_header.add_run("EXPERIENCE")
    exp_run.bold = True
    exp_run.font.size = Pt(12)
    exp_run.font.color.rgb = GREEN
    add_bottom_border(exp_header)
    set_paragraph_spacing(exp_header, before=200, after=100)

    MONTHS = ("", "Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec")

    def fmt_date(d):
        if d == "present":
            return "Present"
        if len(d) >= 7:
            return f"{MONTHS[int(d[5:7])]} {d[:4]}"
        return d[:4]

    def fmt_year(d):
        if d == "present":
            return "Present"
        return d[:4]

    def short_role_title(title: str, all_titles: list[str]) -> str:
        # Compact titles for promotion ladders like "X Applied Scientist"
        suffix = " Applied Scientist"
        if all(t.endswith(suffix) for t in all_titles):
            return title.removesuffix(suffix)
        return title

    for job in filter_included(profile.get("experience", []), target):
        if "roles" in job:
            roles = filter_included(job.get("roles", []), target)
            if not roles:
                continue

            current = next((r for r in roles if r.get("end_date") == "present"), roles[-1])
            prev = [r for r in roles if r is not current]

            all_titles = [r.get("title", "") for r in roles]

            # Header: current title – company
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f'{current["title"]} – {job["company"]}')
            title_run.bold = True
            title_run.font.color.rgb = DARKGREEN
            set_paragraph_spacing(title_para, before=120, after=0)

            # Meta: current dates + location
            meta_para = doc.add_paragraph()
            meta_run = meta_para.add_run(
                f'{fmt_year(current["start_date"])} – {fmt_year(current["end_date"])}, {job["location"]}'
            )
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = LIGHT_GRAY
            set_paragraph_spacing(meta_para, after=20)

            # Previously: compact role list (most recent first)
            if prev:
                prev_sorted = list(reversed(prev))
                prev_bits = []
                for r in prev_sorted:
                    t = short_role_title(r.get("title", ""), all_titles)
                    prev_bits.append(f'{t} ({fmt_year(r["start_date"])}–{fmt_year(r["end_date"])})')
                prev_para = doc.add_paragraph()
                prev_run = prev_para.add_run("Previously: " + " • ".join(prev_bits))
                prev_run.font.size = Pt(10)
                prev_run.font.color.rgb = LIGHT_GRAY
                set_paragraph_spacing(prev_para, after=20)

            # Highlights
            for highlight in filter_included(job.get("highlights", []), target):
                bullet_para = doc.add_paragraph(highlight.get("text", ""), style="List Bullet")
                set_paragraph_spacing(bullet_para, after=40)
        else:
            start_year = job["start_date"][:4]
            end_year = "Present" if job["end_date"] == "present" else job["end_date"][:4]
            # Job title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f'{job["title"]} – {job["company"]}')
            title_run.bold = True
            title_run.font.color.rgb = DARKGREEN
            set_paragraph_spacing(title_para, before=120, after=0)
            # Date/location
            meta_para = doc.add_paragraph()
            meta_run = meta_para.add_run(f'{start_year} – {end_year}, {job["location"]}')
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = LIGHT_GRAY
            set_paragraph_spacing(meta_para, after=60)
            # Highlights
            for highlight in filter_included(job.get("highlights", []), target):
                bullet_para = doc.add_paragraph(highlight.get("text", ""), style="List Bullet")
                set_paragraph_spacing(bullet_para, after=40)

    # Education section
    edu_header = doc.add_paragraph()
    edu_run = edu_header.add_run("EDUCATION")
    edu_run.bold = True
    edu_run.font.size = Pt(12)
    edu_run.font.color.rgb = GREEN
    add_bottom_border(edu_header)
    set_paragraph_spacing(edu_header, before=200, after=100)

    for edu in filter_included(profile.get("education", []), target):
        # Degree
        deg_para = doc.add_paragraph()
        deg_run = deg_para.add_run(f'{edu["degree"]}, {edu["institution"]}')
        deg_run.bold = True
        deg_run.font.color.rgb = DARKGREEN
        set_paragraph_spacing(deg_para, before=80, after=40)
        
        if "focus" in edu:
            doc.add_paragraph(f'Focus: {edu["focus"]}')
        if "thesis" in edu:
            minor = f'. {edu["minor"]} Minor.' if "minor" in edu else ""
            doc.add_paragraph(f'Thesis: {edu["thesis"]}{minor}')
        if "technologies" in edu:
            doc.add_paragraph(f'Technologies: {", ".join(edu["technologies"])}')

    # Publications section
    pub_header = doc.add_paragraph()
    pub_run = pub_header.add_run("PUBLICATIONS")
    pub_run.bold = True
    pub_run.font.size = Pt(12)
    pub_run.font.color.rgb = GREEN
    add_bottom_border(pub_header)
    set_paragraph_spacing(pub_header, before=200, after=100)

    pubs = filter_included(profile.get("publications", []), target)
    for i, pub in enumerate(pubs, 1):
        authors = ", ".join(pub["authors"])

        pub_para = doc.add_paragraph()
        pub_para.add_run(f"{i}. {authors}, \"")
        if "doi" in pub:
            add_hyperlink(pub_para, pub["title"], f'https://doi.org/{pub["doi"]}')
        else:
            pub_para.add_run(pub["title"])
        pub_para.add_run(f',\" {pub["venue"]}')
        if "volume" in pub:
            pub_para.add_run(f", vol. {pub['volume']}")
        if "issue" in pub:
            pub_para.add_run(f", no. {pub['issue']}")
        if "page" in pub:
            pub_para.add_run(f", p. {pub['page']}")
        if "pages" in pub:
            pub_para.add_run(f", pp. {pub['pages']}")
        pub_para.add_run(f", {pub['year']}.")

        set_paragraph_spacing(pub_para, after=60)

    doc.save(output_path)
    print(f"Generated {output_path}", file=sys.stderr)


if __name__ == "__main__":
    main()
